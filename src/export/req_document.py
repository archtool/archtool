'''
Created on Nov 2, 2013

@author: EHWAAL
'''

import re
import StringIO
import model
from model import PlaneableItem
from model import config
from docutils.core import publish_string, default_description
import time, datetime

import docx


TABLE_COLUMNS = ["Id", "Description", "Priority"]
COL_WIDTH     = [718, 7468, 814]

ENCODING = 'utf-8'

TABLE_HEADER = '''
.. list-table:: 
   :widths: 5 85 10
   :header-rows: 1
   
'''

TABLE_START = '''
   * - '''

TABLE_SEPARATOR = '''
     - '''



REQ_NR = re.compile('([0-9a-zA-Z\.]*)[^\.]')

def cmpRequirements(a, b):
  ''' Compare requirements on their names '''
  na, nb = a.Name, b.Name
  na = REQ_NR.match(na).groups()[0]
  nb = REQ_NR.match(nb).groups()[0]
  
  if na == '' or nb == '':
    return cmp(a.Name, b.Name)
  
  aparts = na.split('.')
  bparts = nb.split('.')
  
  for i in range(min([len(aparts), len(bparts)])):
    # Skip this check if the parts are the same.
    if aparts[i] == bparts[i]:
      continue
    if aparts[i].isnumeric() and bparts[i].isnumeric():
      # Compare as integers
      return cmp(int(aparts[i]), int(bparts[i]))
    # Compare as strings
    return cmp(aparts[i], bparts[i])
  # Compare on length...
  return cmp(len(aparts), len(bparts))
  
class RstRenderer(object):
  def __init__(self):
    self.out = StringIO.StringIO()

  def __str__(self):
    return self.out.getvalue()

  def renderChapterHead(self, item, sep='-', also_above=False):
    name = item.Name.encode(ENCODING)
    lines = ['']
    if also_above:
      lines.append(sep*len(name))
    lines += [name, sep*len(name), '']
    if item.Description:
      lines.append(item.Description.encode(ENCODING))
    print >> self.out,  '\n'.join(lines)


  def renderSectionHead(self, item):
    self.renderChapterHead(item, '`')


  def renderRequirementsTable(self, items):
    result = TABLE_HEADER + TABLE_START + TABLE_SEPARATOR.join(TABLE_COLUMNS)

    for item in items:
      values = [str(getattr(item, k)).encode(ENCODING).replace('\n', '\n       ') for k in TABLE_COLUMNS]
      result += TABLE_START + TABLE_SEPARATOR.join(values)

    print >> self.out, result

  def renderVersion(self, latest_mod):
    fmt = '%d %b %Y, %H:%M:%S'
    now = time.strftime(fmt)
    mod = latest_mod.strftime(fmt) if latest_mod else '---'
    print >> self.out, '\n\nLast modification: %s'%mod
    print >> self.out, '\n\nGenerated on: %s'%now

  def write(self, fname):
    html = publish_string(self.out.getvalue(), writer_name='html')

    if not fname.endswith('.html'):
      fname += '.html'

    with open(fname, 'w') as f:
      f.write(html)


class DocxRenderer(object):
  def __init__(self):
    self.out = docx.Document()

  def __str__(self):
    out = StringIO.StringIO()
    self.out.save(out)
    return ''

  def renderChapterHead(self, item, level=0):
    name = item.Name.encode(ENCODING)
    self.out.add_heading(name, level)
    if item.Description:
      self.out.add_paragraph(item.Description.encode(ENCODING))

  def renderSectionHead(self, item):
    self.renderChapterHead(item, 1)


  def renderRequirementsTable(self, items):
    table = self.out.add_table(rows=len(items), cols=len(TABLE_COLUMNS))
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for c, n, w in zip(hdr_cells, TABLE_COLUMNS, COL_WIDTH):
      c.text = n

    for item in items:
      row_cells = table.add_row().cells
      values = [str(getattr(item, k)).encode(ENCODING).replace('\n', '\n       ') for k in TABLE_COLUMNS]
      for c, t in zip(row_cells, values):
        c.text = t

    table.columns[0].width = 1
    table.columns[1].width = 4

  def renderVersion(self, latest_mod):
    fmt = '%d %b %Y, %H:%M:%S'
    now = time.strftime(fmt)
    mod = latest_mod.strftime(fmt) if latest_mod else '---'
    self.out.add_paragraph('Last modification: %s'%mod +
                           '\nGenerated on: %s'%now)

  def write(self, fname):
    if not fname.endswith('.docx'):
      fname += '.docx'

    self.out.save(fname)


def exportRequirementsDocument(session, requirement_name):
  top_items = session.query(model.Requirement).\
                     filter(model.Requirement.Name==requirement_name).\
                     all()

  if len(top_items) == 0:
    raise RuntimeError('Could not find %s, sorry'%requirement_name)
  top_item = top_items[0]

  #renderer = RstRenderer()
  renderer = DocxRenderer()

  last_change = session.query(model.ChangeLog.TimeStamp).\
                       filter(model.ChangeLog.RecordType == model.Requirement.__tablename__).\
                       order_by(model.ChangeLog.TimeStamp.desc()).first()
  
  renderer.renderChapterHead(top_item)
  renderer.renderVersion(last_change[0] if last_change else None)
  for chapter in sorted(top_item.Children, cmp=cmpRequirements):
    renderer.renderChapterHead(chapter)
    
    for section in sorted(chapter.Children, cmp=cmpRequirements):
      renderer.renderSectionHead(section)
      # TODO: make this a recursive query to catch sub-requirements
      #renderer.renderRequirementsTable(sorted(section.Children, key=lambda x:x.Id))
      renderer.renderRequirementsTable(sorted(section.Children, cmp=cmpRequirements))

  print renderer
  fname = '%s/%s'%(config.getConfig('export_dir'), top_item.Name)
  renderer.write(fname)